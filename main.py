import PyPDF2
import openpyxl
import docx

def pdf_read():
    # Open the PDF file
    pdf_ch = open('chi.pdf', 'rb')
    pdf_en = open('eng.pdf', 'rb')

    paragraphs_ch = []
    paragraphs_en = []

    # Create a PDF reader object
    pdf_reader = PyPDF2.PdfReader(pdf_ch)

    # Loop through each page in the PDF document
    for page_num in range(len(pdf_reader.pages)):
        # Get the page object
        page = pdf_reader.pages[page_num]

        # Extract the text from the page
        text = page.extract_text()

        # Split the text into paragraphs
        paragraphs_ch = text.split('\n\n')

    # Create a PDF reader object
    pdf_reader = PyPDF2.PdfReader(pdf_en)

    # Loop through each page in the PDF document
    for page_num in range(len(pdf_reader.pages)):
        # Get the page object
        page = pdf_reader.pages[page_num]

        # Extract the text from the page
        text = page.extract_text()

        # Split the text into paragraphs
        paragraphs_en = text.split('\n\n')

if __name__ == "__main__":

    paragraphs_ch = []
    paragraphs_en = []

    # Open the document
    doc = docx.Document('chi.docx')

    # Loop through each paragraph in the document
    for para in doc.paragraphs:
        # Do something with the paragraph
        txt = para.text.strip()
        if len(txt) and not (txt.startswith('-') and txt.endswith('-')):
            paragraphs_ch.append(para.text.strip())
    
    # Open the document
    doc = docx.Document('eng.docx')

    # Loop through each paragraph in the document
    for para in doc.paragraphs:
        # Do something with the paragraph
        txt = para.text.strip()
        if len(txt) and not (txt.startswith('-') and txt.endswith('-')):
            paragraphs_en.append(para.text.strip())

    paragraphs_en[1] = paragraphs_en[0] + paragraphs_en[1]  
    # Create a new workbook object
    workbook = openpyxl.Workbook()

    # Select the active worksheet
    worksheet = workbook.active
    # Add data to the worksheet
    worksheet['A1'] = 'Eng'
    worksheet['B1'] = 'Chi'
    
    # if len(paragraphs_en) == len(paragraphs_ch):
    if True:
        for i in range(len(paragraphs_ch)):
            # Set the value of cell B2
            worksheet.cell(row=i+2, column=1).value = paragraphs_en[i+1]
            worksheet.cell(row=i+2, column=2).value = paragraphs_ch[i]

        # Save the workbook to a file
        workbook.save('result.xlsx')
